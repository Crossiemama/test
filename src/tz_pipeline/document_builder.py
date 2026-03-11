from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _setup_gost_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _add_paragraph(doc: Document, text: str) -> None:
    stripped = text.strip()
    if not stripped:
        doc.add_paragraph("")
        return

    if stripped.startswith("# "):
        p = doc.add_paragraph(stripped[2:])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        return

    if stripped.startswith("## ") or stripped.startswith("### "):
        title = stripped.replace("## ", "").replace("### ", "")
        p = doc.add_paragraph(title)
        p.runs[0].bold = True
        p.paragraph_format.first_line_indent = Cm(0)
        return

    p = doc.add_paragraph(stripped)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


def markdown_to_gost_docx(markdown_text: str, output_path: str) -> Path:
    doc = Document()
    _setup_gost_layout(doc)

    for line in markdown_text.splitlines():
        if line.strip() == "---":
            continue
        _add_paragraph(doc, line)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output


def convert_markdown_file_to_docx(markdown_path: str, output_path: str) -> Path:
    text = Path(markdown_path).read_text(encoding="utf-8")
    return markdown_to_gost_docx(text, output_path)
